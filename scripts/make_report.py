"""배포된 사이트 스크린샷 캡처 후 Word 과제 문서 생성."""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright

BASE_URL = "https://newsboard-two.vercel.app"

PAGES = [
    ("대시보드 (Overview)", "/", "전체 지표를 한눈에 확인하는 메인 화면"),
    ("주요 이슈", "/issue", "클러스터링으로 묶인 주요 뉴스 이슈 목록"),
    ("경쟁사 비교", "/compare", "인기 랭킹 · 섹션별 랭킹 기반 경쟁사 비교"),
    ("미보도 탐지 (Gap Detection)", "/gap", "경쟁사가 보도했지만 자사가 놓친 이슈 목록"),
    ("구독자 분석", "/analytics/subscribers", "매체별 네이버 구독자 추이 분석"),
    ("독자 반응 (댓글 분석)", "/analytics/comments", "댓글·좋아요 기준 인기 기사 분석"),
    ("AI 리포트", "/report", "AI가 생성한 일간 브리핑 및 이슈 요약"),
]

FEATURES = [
    {
        "title": "대시보드 (Overview)",
        "desc": (
            "자사 오늘 기사 수, 구독자 수, 일일 증감, 댓글 반응 등 핵심 지표를 StatCard로 요약 표시합니다. "
            "AI 일간 브리핑, 주요 이슈 카드, 최신 랭킹 기사, 미보도 알림, 구독자 차트, 인기 댓글 기사를 "
            "단일 화면에서 제공합니다."
        ),
        "stack": [
            "Next.js 15 App Router — Server Component로 Supabase 직접 조회, force-dynamic 렌더링",
            "Supabase JS — daily_publication_count, comment_metric, subscriber_snapshot, issue_cluster, missed_issue_alert 다중 테이블 병렬 조회",
            "Recharts — 구독자 추이 선형 차트",
            "lucide-react — 아이콘",
            "Tailwind CSS — 반응형 그리드 레이아웃 (StatCard 4열, IssueCard 4열)",
        ],
    },
    {
        "title": "주요 이슈",
        "desc": (
            "AI 임베딩 기반 클러스터링으로 묶인 뉴스 이슈를 카드 형태로 표시합니다. "
            "이슈 제목, 요약, 키워드, 참여 매체 수, 신뢰도 점수를 보여주며, "
            "카드 클릭 시 이슈 상세 페이지로 이동해 관련 기사 전체 목록과 AI 요약을 확인할 수 있습니다."
        ),
        "stack": [
            "Next.js 15 App Router — 동적 라우트 /issue/[cluster_id]",
            "Supabase JS — issue_cluster + issue_cluster_article + article 조인 조회",
            "OpenAI text-embedding-3-small — 기사 제목 임베딩 생성",
            "Python 그리디 클러스터링 (scripts/cluster_articles.py) — cosine 유사도 기반 threshold=0.85",
            "OpenAI gpt-4o-mini — 클러스터 대표 제목·요약·키워드 AI 생성 (generate_cluster_metadata)",
            "GitHub Actions cron-cluster.yml — ranking 성공 직후 자동 실행",
        ],
    },
    {
        "title": "경쟁사 비교",
        "desc": (
            "조선·중앙·동아·매경 등 경쟁 매체와 세계일보의 인기 랭킹 및 섹션별 랭킹을 비교합니다. "
            "언론사 칩으로 비교 대상을 선택할 수 있으며, 세계일보는 항상 고정 강조 표시됩니다. "
            "인기 랭킹 탭과 섹션별 랭킹 탭으로 분리되어 있습니다."
        ),
        "stack": [
            "Next.js 15 App Router + 'use client' — 탭·체크박스 상태 관리",
            "Supabase JS — ranking_news_snapshot + ranking_news_item + media_company 조인",
            "Python BeautifulSoup + httpx — 네이버 media.naver.com 랭킹 페이지 스크래핑",
            "GitHub Actions cron-ranking.yml — 매시 UTC :07 자동 수집",
            "GitHub Actions cron-section-ranking.yml — ranking 성공 직후 섹션별 수집",
        ],
    },
    {
        "title": "미보도 탐지 (Gap Detection)",
        "desc": (
            "경쟁사 2개 이상 매체가 보도했지만 세계일보가 보도하지 않은 이슈를 자동 탐지합니다. "
            "우선순위(높음·보통·낮음)별로 색상 구분하여 표시하고, "
            "'검토 시작' 버튼으로 open → reviewing, '완료' 버튼으로 reviewing → resolved 상태를 관리합니다."
        ),
        "stack": [
            "Python scripts/detect_gap.py — issue_cluster 기반 미보도 이슈 탐지, priority_score 계산",
            "Next.js Server Actions (gap/actions.ts) — markReviewing / markResolved, revalidatePath로 즉시 UI 갱신",
            "'use client' ReviewButton + useTransition — 비동기 상태 버튼",
            "Supabase JS — missed_issue_alert CRUD",
            "GitHub Actions cron-gap.yml — cluster 성공 직후 자동 실행",
        ],
    },
    {
        "title": "구독자 분석",
        "desc": (
            "네이버 뉴스 매체별 구독자 수 추이를 표와 차트로 분석합니다. "
            "체크박스로 비교 대상 매체를 선택하고, 구독자 수·일일 증감 토글로 관점을 전환할 수 있습니다. "
            "세계일보는 항상 고정 표시됩니다."
        ),
        "stack": [
            "Next.js 15 App Router + 'use client' — 체크박스·토글 상호작용",
            "Recharts LineChart — 다중 계열 구독자 추이 시각화",
            "Supabase JS — subscriber_snapshot + media_company 조인, 최근 15일 시계열 조회",
            "Python scripts/collect_subscribers.py — 네이버 followers.json API 파싱",
            "GitHub Actions cron-subscribers.yml — UTC 23:00 (KST 08:00) 일 1회 수집",
        ],
    },
    {
        "title": "독자 반응 (댓글 분석)",
        "desc": (
            "댓글 수 기준 인기 기사를 자사(세계일보)와 경쟁사(조선·중앙·동아·매경)로 분리하여 표시합니다. "
            "500개 이상 매우 활발, 200개 이상 활발, 미만 보통으로 참여도 배지를 표시합니다. "
            "동일 기사의 시간별 중복 스냅샷은 최신 1건으로 자동 제거합니다."
        ),
        "stack": [
            "Next.js 15 App Router — Server Component, getOurTopComments / getCompetitorTopComments 쿼리",
            "Supabase JS — comment_metric + article + media_company 조인, article_id 기준 중복 제거",
            "Python scripts/collect_comments.py + Playwright — 네이버 랭킹 페이지 댓글 수 스크래핑",
            "GitHub Actions cron-comments.yml — 매시 UTC :15 (KST :24) 수집, 네이버 업데이트(:15) 9분 후",
        ],
    },
    {
        "title": "AI 리포트",
        "desc": (
            "OpenAI gpt-4o-mini가 생성한 일간 뉴스 브리핑과 이슈별 심층 요약을 제공합니다. "
            "일간 브리핑은 매일 KST 00:00에 자동 생성되며, "
            "이슈 상세 페이지에서 개별 이슈에 대한 AI 요약을 온디맨드로 생성할 수 있습니다."
        ),
        "stack": [
            "FastAPI (Vercel Fluid Compute) — /api/report/daily, /api/report/issue/{id} 엔드포인트",
            "OpenAI gpt-4o-mini — 브리핑 제목·요약·핵심 포인트 JSON 생성",
            "Supabase JS — ai_summary 테이블 upsert (summary_type, summary_date 복합 키)",
            "GitHub Actions cron-daily-briefing.yml — UTC 15:00 (KST 00:00) 자동 실행",
            "Next.js Server Component — /report 페이지, AI 요약 카드 렌더링",
        ],
    },
]


def capture_screenshots(tmp_dir: Path) -> dict[str, Path]:
    paths: dict[str, Path] = {}
    print("스크린샷 캡처 시작...")
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 1440, "height": 900})
        for title, path, _ in PAGES:
            url = BASE_URL + path
            print(f"  캡처: {title} ({url})")
            try:
                page.goto(url, wait_until="networkidle", timeout=30000)
                page.wait_for_timeout(1500)
                img_path = tmp_dir / f"{title.replace('/', '_').replace(' ', '_')}.png"
                page.screenshot(path=str(img_path), full_page=False)
                paths[title] = img_path
            except Exception as e:
                print(f"  ✗ {title} 캡처 실패: {e}")
        browser.close()
    return paths


def build_docx(screenshots: dict[str, Path], out_path: Path) -> None:
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # 표지
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Newsboard")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_para.add_run("AI 기반 미디어 모니터링 대시보드")
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    overview_para = doc.add_paragraph()
    overview_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ov = overview_para.add_run(
        "뉴스 조직 내부용 AI 미디어 모니터링 및 인사이트 대시보드.\n"
        "자사(세계일보)와 경쟁사 기사를 자동 수집·클러스터링하여\n"
        "미보도 탐지, 경쟁사 비교, 독자 반응 분석 기능을 제공합니다."
    )
    ov.font.size = Pt(11)

    doc.add_page_break()

    # 기술 스택 개요
    h = doc.add_heading("전체 기술 스택", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    stack_table = doc.add_table(rows=1, cols=2)
    stack_table.style = "Table Grid"
    hdr = stack_table.rows[0].cells
    hdr[0].text = "영역"
    hdr[1].text = "기술"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    stack_rows = [
        ("프론트엔드", "Next.js 15 App Router, Tailwind CSS, Recharts, lucide-react"),
        ("백엔드 API", "FastAPI (Python 3.13, Vercel Fluid Compute)"),
        ("데이터베이스", "Supabase (PostgreSQL) — 11개 테이블, 3개 마이그레이션"),
        ("AI", "OpenAI gpt-4o-mini (요약·메타 생성), text-embedding-3-small (클러스터링)"),
        ("데이터 수집", "Python httpx + BeautifulSoup, Playwright (JavaScript 렌더링)"),
        ("자동화", "GitHub Actions 8종 cron 워크플로"),
        ("배포", "Vercel (프론트 + Python API 단일 프로젝트)"),
        ("언어", "TypeScript (프론트), Python (스크립트·API)"),
    ]
    for left, right in stack_rows:
        row = stack_table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_paragraph()
    doc.add_page_break()

    # 기능별 상세
    h2 = doc.add_heading("기능 상세", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    for feat in FEATURES:
        # 기능 제목
        fh = doc.add_heading(feat["title"], level=2)
        fh.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        # 기능 설명
        dp = doc.add_paragraph(feat["desc"])
        dp.paragraph_format.space_after = Pt(6)

        # 기술 스택
        sp = doc.add_paragraph()
        sp.add_run("사용 기술").bold = True
        for item in feat["stack"]:
            doc.add_paragraph(item, style="List Bullet")

        # 스크린샷
        if feat["title"] in screenshots:
            doc.add_paragraph()
            doc.add_picture(str(screenshots[feat["title"]]), width=Inches(6.0))
            cap = doc.add_paragraph(f"▲ {feat['title']} 화면")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        else:
            doc.add_paragraph("(스크린샷 캡처 실패)").runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        doc.add_paragraph()

    doc.save(str(out_path))
    print("\nDone: " + str(out_path))


def main() -> None:
    out_path = Path("d:/newsboard/Newsboard_과제보고서.docx")
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        screenshots = capture_screenshots(tmp_dir)
        print(f"\n캡처 완료: {len(screenshots)}/{len(PAGES)}개")
        build_docx(screenshots, out_path)


if __name__ == "__main__":
    main()
